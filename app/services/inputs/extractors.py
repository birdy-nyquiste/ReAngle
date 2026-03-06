"""
根据不同输入源提取文本
"""

import httpx
import docx
import pypdf
import io
from bs4 import BeautifulSoup
from readability import Document
from fastapi import UploadFile
from loguru import logger

from app.core.exceptions import ContentExtractionError

from PIL import Image

# 轻量 OCR（中英混排效果较好）
try:
    from rapidocr_onnxruntime import RapidOCR  # type: ignore
    import numpy as np  # 仅在 RapidOCR 可用时再导入 numpy
    _rapid_ocr = RapidOCR()  # 初始化一次复用
except Exception:
    _rapid_ocr = None
    np = None  # type: ignore
    logger.warning("RapidOCR is not available. Image OCR will be disabled.")

async def extract_text_from_url(url: str) -> str:
    """
    从 URL 提取主要文本内容。
    """
    try:
        logger.info(f"Starting URL extraction: {url}")

        # 确保URL有协议
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
            logger.debug(f"URL protocol added: {url}")

        async with httpx.AsyncClient() as client:
            logger.debug("Sending HTTP request...")
            response = await client.get(url, timeout=30.0)
            logger.debug(f"HTTP response status: {response.status_code}")
            response.raise_for_status()

        logger.debug("Parsing HTML content...")
        doc = Document(response.text)
        summary = doc.summary()
        logger.debug(f"Readability summary length: {len(summary)}")

        soup = BeautifulSoup(summary, "html.parser")
        text = soup.get_text()

        # 清理文本
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        result = "\n".join(lines)
        logger.info(f"Extraction complete. Length: {len(result)}")

        if not result:
            raise ContentExtractionError("Extracted text is empty")

        return result

    except httpx.HTTPError as e:
        logger.error(f"HTTP error during extraction: {e}")
        raise ContentExtractionError(f"Failed to download URL: {str(e)}")
    except Exception as e:
        logger.exception(f"Unexpected error extracting from URL: {url}")
        raise ContentExtractionError(f"Error extracting from URL: {str(e)}")


async def extract_text_from_docx(file: UploadFile) -> str:
    """
    从 DOCX 文件提取文本。
    """
    try:
        logger.info(f"Starting DOCX extraction: {file.filename}")
        content = await file.read()
        doc = docx.Document(io.BytesIO(content))

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        result = "\n".join(text_parts)
        logger.info(f"DOCX extraction complete. Length: {len(result)}")
        return result

    except Exception as e:
        logger.exception(f"Error extracting from DOCX: {file.filename}")
        raise ContentExtractionError(f"Error extracting from DOCX: {str(e)}")


async def extract_text_from_pdf(file: UploadFile) -> str:
    """
    从 PDF 文件提取文本。
    """
    try:
        logger.info(f"Starting PDF extraction: {file.filename}")
        content = await file.read()
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))

        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())

        result = "\n".join(text_parts)
        logger.info(f"PDF extraction complete. Length: {len(result)}")
        return result

    except Exception as e:
        logger.exception(f"Error extracting from PDF: {file.filename}")
        raise ContentExtractionError(f"Error extracting from PDF: {str(e)}")


async def extract_text_from_image(file: UploadFile) -> str:
    """
    从图片文件提取文本（OCR）。支持常见格式：png/jpg/jpeg/webp。
    优先使用 RapidOCR；若不可用则报错提示安装。
    """
    if _rapid_ocr is None:
        raise ContentExtractionError(
            "OCR not available. Please install rapidocr-onnxruntime."
        )
    try:
        logger.info(f"Starting IMAGE OCR: {file.filename}")
        content = await file.read()
        # 基础预处理：转 RGB，尽量消除模式差异
        image = Image.open(io.BytesIO(content)).convert("RGB")
        np_img = np.array(image)
        result, _ = _rapid_ocr(np_img)
        lines = []
        for item in (result or []):
            # 兼容不同版本 RapidOCR 的返回：
            # 1) [box, text, score]
            # 2) [box, (text, score)]
            text = ""
            if isinstance(item, (list, tuple)):
                if len(item) >= 3:
                    # [box, text, score]
                    text = str(item[1] or "")
                elif len(item) == 2:
                    second = item[1]
                    if isinstance(second, (list, tuple)) and second:
                        text = str(second[0] or "")
                    else:
                        text = str(second or "")
            else:
                # 非预期结构，跳过
                continue
            if text.strip():
                lines.append(text.strip())
        extracted = "\n".join(lines)
        logger.info(f"IMAGE OCR complete. Length: {len(extracted)}")
        if not extracted.strip():
            # 给到用户可读的错误，便于前端提示
            raise ContentExtractionError("OCR result is empty. Try a clearer image.")
        return extracted
    except ContentExtractionError:
        raise
    except Exception as e:
        logger.exception(f"Error extracting from IMAGE: {file.filename}")
        raise ContentExtractionError(f"Error extracting from IMAGE: {str(e)}")
